from pathlib import Path
from zipfile import ZipFile
from lxml import etree
from docx import Document


path = Path(__file__).resolve().parents[1] / "docs" / "AD3306_AD3304_LED_Communication_Protocol_OEM_KR.docx"
doc = Document(path)

assert len(doc.paragraphs) > 100, "본문 단락 수가 예상보다 적음"
assert len(doc.tables) >= 18, "표 수가 예상보다 적음"

text = "\n".join(p.text for p in doc.paragraphs)
for required in (
    "AD3306-AD3304 LED 통신 프로토콜 사양서",
    "OspBridge_SetColorRgbByNodes",
    "FIO-OSP SET_PWM 프레임",
    "A0 17 4F FF FF C0 40 80 00 64",
    "OEM 통합 요구사항",
    "OEM 검증 절차",
):
    assert required in text or any(required in c.text for t in doc.tables for r in t.rows for c in r.cells), required

for idx, table in enumerate(doc.tables, 1):
    cols = len(table.rows[0].cells)
    assert all(len(row.cells) == cols for row in table.rows), f"table {idx}: column mismatch"
    assert len(table.rows) >= 1, f"table {idx}: empty"

with ZipFile(path) as zf:
    xml = etree.fromstring(zf.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tables = xml.xpath("//w:tbl", namespaces=ns)
    assert len(tables) == len(doc.tables)
    for idx, table in enumerate(tables, 1):
        grid = table.xpath("./w:tblGrid/w:gridCol/@w:w", namespaces=ns)
        widths = table.xpath("./w:tr[1]/w:tc/w:tcPr/w:tcW/@w:w", namespaces=ns)
        assert grid, f"table {idx}: missing grid"
        assert widths, f"table {idx}: missing tcW"
        assert grid == widths, f"table {idx}: grid/tcW mismatch {grid} != {widths}"
        assert sum(map(int, grid)) == 9360, f"table {idx}: width != 9360"

print(f"OK paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} size={path.stat().st_size}")
