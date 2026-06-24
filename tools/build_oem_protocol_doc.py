from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path(__file__).resolve().parents[1] / "docs" / "AD3306_AD3304_LED_Communication_Protocol_OEM_KR.docx"

BLUE = "1F4E79"
DARK = "17365D"
LIGHT = "D9EAF7"
PALE = "EEF5FA"
GRAY = "F2F4F7"
MID = "667085"
GREEN = "E2F0D9"
GOLD = "FFF2CC"
RED = "FCE4D6"
WHITE = "FFFFFF"
BLACK = "202124"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="AAB7C4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn("w:" + edge))
        if tag is None:
            tag = OxmlElement("w:" + edge)
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Malgun Gothic", size=10.5, bold=False, color=BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(p, text, bold=False, color=BLACK, size=10.5, code=False):
    run = p.add_run(text)
    set_run_font(run, "Consolas" if code else "Malgun Gothic", size, bold, color)
    return run


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.34)
    sec.footer_distance = Inches(0.34)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after, color in (
        ("Title", 24, 0, 5, DARK),
        ("Subtitle", 11, 0, 12, MID),
        ("Heading 1", 16, 15, 8, BLUE),
        ("Heading 2", 13, 11, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, DARK),
    ):
        st = styles[name]
        st.font.name = "Malgun Gothic"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "AD3306-AD3304 LED Communication Protocol | OEM", False, MID, 8.5)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "CONFIDENTIAL - Engineering Reference   |   ", False, MID, 8)
    add_field(footer, "PAGE")
    add_text(footer, " / ", False, MID, 8)
    add_field(footer, "NUMPAGES")


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=8.9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, h, True, DARK, font_size)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, str(value), False, BLACK, font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, title, body, fill=PALE):
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t, color=BLUE, size="8")
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_text(p, title + "  ", True, DARK, 10)
    add_text(p, body, False, BLACK, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.38 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    add_text(p, text, False, BLACK, 10.2)
    return p


def add_step(doc, number, title, body):
    t = doc.add_table(rows=1, cols=2)
    set_table_borders(t, color="C5D3E0", size="6")
    set_table_geometry(t, [720, 8640])
    c0, c1 = t.rows[0].cells
    set_cell_shading(c0, BLUE)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p0, str(number), True, WHITE, 14)
    p1 = c1.paragraphs[0]
    add_text(p1, title, True, DARK, 10.5)
    p1.add_run("\n")
    add_text(p1, body, False, BLACK, 9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        add_text(p, line, False, BLACK, 8.5, code=True)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    style_doc(doc)

    # Cover / masthead
    p = doc.add_paragraph(style="Title")
    add_text(p, "AD3306-AD3304 LED 통신 프로토콜 사양서", True, DARK, 24)
    p = doc.add_paragraph(style="Subtitle")
    add_text(p, "FT4222(SPI) - AD3306(10BASE-T1S Controller) - AD3304(Flex IO/FIO-OSP) - RGB LED", False, MID, 11)

    add_table(doc, ["문서 항목", "내용"], [
        ("문서 목적", "OEM 애플리케이션에서 LED 제어 데이터를 생성하여 실제 LED PWM이 변경되기까지의 종단 간 통신 경로 정의"),
        ("적용 프로젝트", "AD3304_20260611_LED25_Up"),
        ("기준 소스", "FT4222Manager.cs / NativeOsp ex_fioOsp.c / utils fioOsp.c / adi_Eth10BaseT1s_cfg.*"),
        ("문서 버전", "Rev. 1.0"),
        ("작성 기준일", "2026-06-12"),
        ("대상 독자", "OEM SW/HW 개발, 시스템 통합, 검증 및 품질 담당자"),
    ], [1800, 7560], font_size=9.3)

    add_callout(doc, "핵심 제어 계약", "OEM은 FID, LED 체인 주소(nodeAddr), R/G/B(0~255)를 지정한다. OspBridge가 AD3304 선택, RGB-PWM 변환, FIO-OSP SET_PWM 프레임 구성, 10BASE-T1S 전송 및 최종 flush를 수행한다.", GREEN)

    doc.add_heading("문서 범위", level=1)
    add_bullet(doc, "호스트 PC/C# 애플리케이션에서 호출해야 하는 공개 함수와 인자 정의")
    add_bullet(doc, "FT4222-AD3306 SPI 구간과 AD3306-AD3304 10BASE-T1S 구간의 계층별 역할")
    add_bullet(doc, "AD3304 Flex IO를 통한 FIO-OSP LED 초기화 및 SET_PWM 프레임 정의")
    add_bullet(doc, "단일 LED, 복수 LED batch, AD3304 단위 broadcast, 전체 AD3304 multicast 사용 절차")
    add_bullet(doc, "진단 상태, 오류 코드 처리 및 OEM 통합 검증 항목")

    page_break(doc)
    doc.add_heading("1. 시스템 구성 및 책임 경계", level=1)
    add_table(doc, ["계층", "구성 요소", "인터페이스", "주요 책임"], [
        ("Application", "Windows C# FormCIE", "Managed API", "색상/CIE, FID, SlaveFlag, 애니메이션 상태 생성"),
        ("Host Driver", "FT4222Manager.cs", "P/Invoke", "입력 검증, OspBridge DLL 호출, 오류 문자열 전달"),
        ("Native Bridge", "OspBridge.dll", "C ABI", "FID 핸들 선택, RGB-PWM 변환, FIO-OSP 요청 생성"),
        ("Host Physical", "FT4222", "USB-to-SPI", "PC에서 AD3306까지 SPI 전송. 세부 SPI transaction은 ADI PAL/OASPI 스택이 관리"),
        ("Controller", "AD3306", "10BASE-T1S", "PLCA controller, remote discovery, E2B/FIO 메시지 전달"),
        ("Remote", "AD3304", "Flex IO", "수신 FIO 데이터를 로컬 FIO-OSP 인터페이스로 라우팅"),
        ("LED Bus", "RGB LED chain", "OSP", "주소 지정 SET_PWM 수신 후 R/G/B PWM 갱신"),
    ], [1250, 1750, 1500, 4860], font_size=8.5)

    doc.add_heading("1.1 종단 간 데이터 흐름", level=2)
    add_table(doc, ["OEM Application", "Host/Bridge", "Vehicle-side Link", "LED Output"], [
        ("FID + nodeAddr + RGB", "FT4222Manager → OspBridge.dll", "FT4222 SPI → AD3306 → 10BASE-T1S → AD3304 Flex IO", "FIO-OSP SET_PWM → 15-bit RGB PWM"),
    ], [2100, 2350, 3260, 1650], header_fill=BLUE, font_size=8.7)
    add_callout(doc, "책임 경계", "애플리케이션은 SPI 바이트나 Ethernet frame을 직접 만들지 않는다. OEM 제어 인터페이스는 OspBridge 공개 API이며, 하위 transport encapsulation은 ADI EAL/Network/OASPI 스택의 책임이다.", GOLD)

    doc.add_heading("1.2 주소 체계", level=2)
    add_table(doc, ["주소", "범위", "의미", "예"], [
        ("FID", "0x01~0x08", "10BASE-T1S상 AD3304 선택. 발견 MAC의 마지막 바이트와 동일", "MAC 00:E0:22:FE:70:03 → FID 0x03"),
        ("FID 0x00", "예약 broadcast", "연결된 AD3304 multicast group 전체에 동일 RGB 전달", "전체 remote 동시 점등"),
        ("remoteNum", "0~7", "EAL 내부 remote index. 현재 구현은 FID-1", "FID 0x04 → remoteNum 3"),
        ("PLCA ID", "1~8", "AD3304의 10BASE-T1S PLCA node ID. 현재 구현은 FID와 동일", "FID 0x04 → PLCA 4"),
        ("nodeAddr", "1~1002", "선택된 AD3304 뒤 OSP LED chain의 device address", "nodeAddr 5 → 다섯 번째 LED"),
        ("nodeAddr 0", "broadcast", "선택한 FID의 LED chain 전체", "한 AD3304 뒤 모든 LED 동일 색상"),
    ], [1250, 1100, 4300, 2710], font_size=8.7)

    page_break(doc)
    doc.add_heading("2. 연결 및 초기화 시퀀스", level=1)
    add_step(doc, 1, "C# 연결 시작", "FT4222Manager.Connect()가 Native OspBridge 모드를 선택하고 OspBridge_InitFixed()를 호출한다.")
    add_step(doc, 2, "FT4222/SPI 장치 개방", "OspBridge 내부에서 adi_pal_init(), adi_network_init(), adi_network_openDevice()가 호출된다. ADI_NETWORK_DEV_IDX는 SPI 장치 index이다.")
    add_step(doc, 3, "AD3306 E2B instance 및 PLCA 구성", "adi_eal_createInstance() 후 controller PLCA ID 0, 최대 node count 9(Controller 1 + AD3304 최대 8)로 활성화한다.")
    add_step(doc, 4, "AD3304 Discovery", "10BASE-T1S remote discovery 결과 MAC의 마지막 바이트를 FID로 해석한다. 유효 범위는 0x01~0x08이다.")
    add_step(doc, 5, "FID별 FIO handle 생성", "발견된 각 FID에 대해 remoteNum=FID-1, intfNum=0으로 adi_eal_fioOspOpen()을 수행한다.")
    add_step(doc, 6, "LED chain 초기화", "RESET → INIT_BIDIR → SET_SETUP → GO_ACTIVE 순서로 수행한다. 성공 시 할당 node address를 gnFidNodeAddr[FID-1]에 저장한다.")
    add_step(doc, 7, "Runtime thread 시작", "adi_network_run()을 수행하는 worker thread를 시작하여 queue와 flush가 실제 SPI/10BASE-T1S 송신으로 진행되게 한다.")

    doc.add_heading("2.1 FIO-OSP 초기 설정", level=2)
    add_table(doc, ["항목", "설정값", "비고"], [
        ("Instruction clock", "2.4 MHz", "adi_eal_fioOspConfigureInstructionRAM"),
        ("Queue depth", "60", "연속 unicast/batch TX queue overflow 방지"),
        ("PWM frequency", "586 Hz", "ADI_FIO_OSP_PWM_FREQ_586HZ"),
        ("CRC", "Enable", "SETUP register"),
        ("Temperature update", "2.4 kHz", "ADI_FIO_OSP_TEMP_UPDATE_RATE_2_4_KHZ"),
        ("Overtemperature action", "Sleep enable", "기타 COM/LOS/UV sleep은 비활성"),
        ("Device state", "ACTIVE", "FIO_OSP_GO_ACTIVE_CMD = 0x05"),
    ], [2500, 1900, 4960], font_size=8.8)

    doc.add_heading("2.2 정상 연결 판정", level=2)
    add_bullet(doc, "OspBridge_InitFixed() 반환값이 0이어야 한다.")
    add_bullet(doc, "대상 FID의 Open, Callback, RAM, Reset, Bidir, Setup, State가 0이고 Handle valid가 1이어야 한다.")
    add_bullet(doc, "미발견 FID는 -99(미실행) 상태로 유지될 수 있으며 연결 실패로 간주하지 않는다.")
    add_bullet(doc, "현재 FIO1-only 구성에서 FIO0의 -77은 의도적 skip 상태일 수 있다. OEM 판정은 실제 FID별 상태를 우선한다.")

    page_break(doc)
    doc.add_heading("3. OEM 런타임 제어 API", level=1)
    add_table(doc, ["사용 목적", "C# 함수", "입력", "하위 Native 함수"], [
        ("AD3304/chain 전체 동일 색", "SetColorRgbByFid", "fid, r, g, b", "OspBridge_SetColorRgbByFid"),
        ("단일 LED", "SetColorRgbByNode", "fid, nodeAddr, r, g, b", "OspBridge_SetColorRgbByNode"),
        ("복수 LED 서로 다른 색", "SetColorRgbByNodes", "fid, node[], r[], g[], b[], count", "OspBridge_SetColorRgbByNodes"),
        ("CIE xy 색상", "SetColorCIEByFid", "fid, cx, cy, brightness01", "RGB 변환 후 SetColorRgbByFid"),
        ("상태/온도 readback", "ReadOspFeedback", "fid, nodeAddr", "OspBridge_ReadbackByFid"),
    ], [1550, 2100, 3200, 2510], font_size=8.4)

    doc.add_heading("3.1 권장 API 선택", level=2)
    add_table(doc, ["상황", "권장 방식", "이유"], [
        ("한 LED만 변경", "SetColorRgbByNode", "unicast 1건, 호출 후 즉시 flush"),
        ("여러 LED를 한 Tick에 변경", "SetColorRgbByNodes", "노드별 frame을 queue에 모은 뒤 마지막에 1회 flush"),
        ("선택한 AD3304의 모든 LED 동일 색", "SetColorRgbByFid(fid)", "할당 주소 또는 chain broadcast로 1회 전송"),
        ("모든 AD3304 동일 색", "SetColorRgbByFid(0x00)", "multicast handle로 전체 remote 동시 전송"),
        ("Shift/RGBW animation", "SetColorRgbByNodes", "Tick마다 batch 전송하여 host/native 호출 및 flush 비용 최소화"),
    ], [2300, 2600, 4460], font_size=8.7)

    doc.add_heading("3.2 C# 호출 예", level=2)
    add_code(doc, "// FID 0x03의 LED #5를 Orange 계열로 변경\nbool ok = ft4222.SetColorRgbByNode(\n    fid: 0x03, nodeAddr: 5, r: 255, g: 128, b: 0);\n\n// FID 0x03의 LED #1~#4를 서로 다른 색으로 batch 변경\nushort[] nodes = { 1, 2, 3, 4 };\nbyte[] red   = { 255, 0,   0,   255 };\nbyte[] green = { 0,   255, 0,   255 };\nbyte[] blue  = { 0,   0,   255, 255 };\nbool batchOk = ft4222.SetColorRgbByNodes(\n    0x03, nodes, red, green, blue, nodes.Length);")
    add_callout(doc, "입력 검증", "count는 모든 배열 길이 이하이어야 하며, nodeAddr는 1-based이다. count=0은 성공(no-op)으로 처리된다. FID는 발견되어 유효 handle이 열린 0x01~0x08이어야 한다.", GOLD)

    page_break(doc)
    doc.add_heading("4. 런타임 데이터 변환 및 전송 순서", level=1)
    add_step(doc, 1, "색상 입력", "애플리케이션이 R/G/B 각 0~255 값을 생성한다. CIE 입력인 경우 먼저 CIE xyY에서 RGB byte로 변환한다.")
    add_step(doc, 2, "AD3304 선택", "OspBridge는 fid를 인덱스(fid-1)로 변환하여 gahFidFioOspHandle[]에서 해당 remote FIO handle을 선택한다.")
    add_step(doc, 3, "LED 주소 선택", "nodeAddr=1~1002는 unicast, nodeAddr=0은 선택된 chain broadcast로 변환한다.")
    add_step(doc, 4, "8-bit RGB → 15-bit PWM", "각 채널은 PWM = floor((RGB × 0x7FFF + 127) / 255)로 변환한다. DayMode bit는 1로 설정한다.")
    add_step(doc, 5, "FIO-OSP frame 생성", "adi_fioOspSetPwm()이 Preamble/Address/PSI/CMD/PWM/CRC의 10-byte SET_PWM frame을 만든다.")
    add_step(doc, 6, "EAL queue 적재", "adi_eal_fioOspWrite()가 선택된 AD3304 remote/intf queue에 frame을 적재한다. write-only ID 0xFF를 사용한다.")
    add_step(doc, 7, "Flush 및 물리 전송", "단일 전송은 즉시 flush, batch는 모든 node frame 적재 후 1회 flush한다. ADI Network/OASPI가 FT4222 SPI와 AD3306 전송을 수행한다.")
    add_step(doc, 8, "10BASE-T1S remote 수신", "AD3306이 PLCA 기반 10BASE-T1S 링크를 통해 대상 AD3304로 E2B/FIO payload를 전달한다.")
    add_step(doc, 9, "Flex IO → LED", "AD3304의 FIO0/LCE-IF0 routing이 payload를 LED OSP bus로 출력하고, 대상 LED가 R/G/B PWM register를 갱신한다.")

    doc.add_heading("4.1 RGB-PWM 변환 예", level=2)
    add_table(doc, ["RGB byte", "15-bit PWM", "Hex", "DayMode 포함 상위 byte"], [
        ("0", "0", "0x0000", "0x80"),
        ("128", "16448", "0x4040", "0xC0"),
        ("255", "32767", "0x7FFF", "0xFF"),
    ], [1800, 2000, 1800, 3760], font_size=9)

    doc.add_heading("4.2 Batch 전송 타이밍", level=2)
    add_table(doc, ["구분", "Native 처리", "Flush 횟수", "적용"], [
        ("단일 node", "frame 생성 → write → flush", "호출당 1회", "SetColorRgbByNode"),
        ("N개 batch", "N개 frame 생성 → N회 write → 최종 flush", "batch당 1회", "SetColorRgbByNodes"),
        ("FID broadcast", "1개 broadcast frame → flush", "호출당 1회", "SetColorRgbByFid"),
    ], [1600, 4100, 1400, 2260], font_size=8.8)

    page_break(doc)
    doc.add_heading("5. FIO-OSP SET_PWM 프레임", level=1)
    add_callout(doc, "Frame length", "SET_PWM request는 총 10 byte이다. 응답을 기다리지 않는 write-only 전송은 request ID 0xFF를 사용한다.", PALE)
    add_table(doc, ["Byte", "필드", "비트 정의", "설명"], [
        ("0", "Preamble + Address[9:6]", "0xA0 | ((addr & 0x3C0) >> 6)", "Preamble 0xA0, 상위 주소"),
        ("1", "Address[5:0] + PSI[2:1]", "((addr & 0x3F) << 2) | 0x03", "10-byte frame의 PSI=6"),
        ("2", "PSI[0] + Command", "0x4F", "SET_PWM command. PSI[0]=0"),
        ("3", "Red High + DayMode", "0x80 | R_PWM[14:8]", "DayMode=1"),
        ("4", "Red Low", "R_PWM[7:0]", "Red 15-bit PWM"),
        ("5", "Green High + DayMode", "0x80 | G_PWM[14:8]", "DayMode=1"),
        ("6", "Green Low", "G_PWM[7:0]", "Green 15-bit PWM"),
        ("7", "Blue High + DayMode", "0x80 | B_PWM[14:8]", "DayMode=1"),
        ("8", "Blue Low", "B_PWM[7:0]", "Blue 15-bit PWM"),
        ("9", "CRC", "CRC-8 polynomial 0x2F", "Byte 0~8에 대해 계산"),
    ], [720, 2100, 3100, 3440], font_size=8.35)

    doc.add_heading("5.1 실제 frame 예제", level=2)
    add_table(doc, ["제어 조건", "PWM 변환", "10-byte frame (hex)"], [
        ("nodeAddr=1, RGB=(255,0,0)", "R=0x7FFF, G=0x0000, B=0x0000", "A0 07 4F FF FF 80 00 80 00 C0"),
        ("nodeAddr=5, RGB=(255,128,0)", "R=0x7FFF, G=0x4040, B=0x0000", "A0 17 4F FF FF C0 40 80 00 64"),
        ("broadcast, RGB=(0,255,0)", "R=0x0000, G=0x7FFF, B=0x0000", "A0 03 4F 80 00 FF FF 80 00 92"),
    ], [2500, 2900, 3960], font_size=8.5)
    add_callout(doc, "주의", "위 10-byte 값은 LED-side FIO-OSP payload 기준이다. FT4222 SPI transaction 및 AD3306 10BASE-T1S Ethernet/E2B encapsulation에는 ADI transport header가 추가되며 애플리케이션 API에서 직접 구성하지 않는다.", GOLD)

    page_break(doc)
    doc.add_heading("6. 인터페이스별 상세 동작", level=1)
    doc.add_heading("6.1 FT4222 → AD3306 (SPI)", level=2)
    add_bullet(doc, "OspBridge는 ADI PAL/Network/OASPI 라이브러리를 통해 FT4222 SPI device를 open한다.")
    add_bullet(doc, "애플리케이션은 SPI mode, chip-select transaction, OASPI framing을 직접 제어하지 않는다.")
    add_bullet(doc, "adi_network_flushTxBuffer() 호출 시 queue의 E2B/FIO 요청이 OASPI를 통해 AD3306으로 전달된다.")
    add_bullet(doc, "OEM 교체 구현 시 OspBridge C ABI와 동등한 synchronous status/error contract를 유지해야 한다.")

    doc.add_heading("6.2 AD3306 → AD3304 (10BASE-T1S)", level=2)
    add_table(doc, ["항목", "현재 설정"], [
        ("Controller MAC", "00:E0:22:FE:70:F0"),
        ("Remote MAC base", "00:E0:22:FE:70:00 + FID"),
        ("Controller PLCA ID", "0"),
        ("Remote PLCA ID", "FID 1~8"),
        ("Maximum node count", "9"),
        ("Remote selection", "발견 MAC 마지막 byte → FID → remoteNum(FID-1)"),
    ], [3000, 6360], font_size=9)

    doc.add_heading("6.3 AD3304 → LED (Flex IO/FIO-OSP)", level=2)
    add_bullet(doc, "Remote register map에서 FIO0은 LCE IF0으로, FIO1은 LCE IF1으로 routing된다.")
    add_bullet(doc, "현재 FID별 handle 설정은 nIntfNum=0을 사용하며 AD3304의 FIO path를 통해 OSP frame을 출력한다.")
    add_bullet(doc, "INIT_BIDIR 후 할당된 LED address를 저장하고 이후 SET_PWM unicast target으로 사용한다.")
    add_bullet(doc, "nodeAddr=0은 chain broadcast이고, nodeAddr=1 이상은 개별 LED address이다.")

    doc.add_heading("6.4 Broadcast 계층 구분", level=2)
    add_table(doc, ["호출", "범위", "사용 handle/주소", "결과"], [
        ("SetColorRgbByFid(0x00,...)", "모든 연결 AD3304", "multicast handle + LED broadcast", "모든 remote/LED 동일 RGB"),
        ("SetColorRgbByFid(FID,...)", "선택 AD3304", "FID handle + assigned/broadcast target", "선택 remote의 chain 변경"),
        ("SetColorRgbByNode(FID,0,...)", "선택 AD3304", "FID handle + address 0", "선택 remote의 모든 LED"),
        ("SetColorRgbByNode(FID,N,...)", "LED 1개", "FID handle + address N", "해당 LED만 변경"),
    ], [2600, 1900, 2600, 2260], font_size=8.5)

    page_break(doc)
    doc.add_heading("7. 기능별 송신 방식", level=1)
    add_table(doc, ["기능", "데이터 생성", "송신 API", "특징"], [
        ("일반 CIE/RGB", "선택 LED 또는 전체에 동일 RGB", "ByFid 또는 ByNode(s)", "선택 범위에 따라 broadcast/unicast"),
        ("RGBW Shift", "node index와 phase에 따라 R/G/B/W 배열 생성", "SetColorRgbByNodes", "한 Tick의 모든 node frame 후 1회 flush"),
        ("Shift", "이동 LED 주변 밝기 wave level 계산", "SetColorRgbByNodes", "변경된 level node만 batch"),
        ("Shift Stack", "이동 LED + 끝점 locked LED", "SetColorRgbByNodes", "누적 상태 포함 batch"),
        ("Shift Up", "중앙에서 양쪽으로 이동 node 계산", "SetColorRgbByNodes", "양방향 wave"),
        ("Shift Stack Up", "양쪽 이동 + 끝점 inward locked", "SetColorRgbByNodes", "양쪽 누적 상태 포함"),
    ], [1600, 3200, 2300, 2260], font_size=8.3)
    add_callout(doc, "성능 원칙", "애니메이션 Tick에서 노드별 DLL 호출과 flush를 반복하지 않는다. node/R/G/B 배열을 먼저 구성하고 SetColorRgbByNodes() 1회로 전달한다.", GREEN)

    doc.add_heading("7.1 Batch 실패 처리", level=2)
    add_bullet(doc, "배열 null, 길이 불일치, count 초과는 C# layer에서 즉시 실패한다.")
    add_bullet(doc, "Native loop 중 한 node가 실패하면 해당 status를 반환하고 이후 node는 전송하지 않는다.")
    add_bullet(doc, "성공한 node frame이 이미 queue에 적재되었을 수 있으므로, OEM safety 요구 시 상위에서 전체 상태 재동기화 frame을 재전송한다.")
    add_bullet(doc, "오류 상세는 FT4222Manager.LastError 및 OspBridge_GetLastError()로 확인한다.")

    doc.add_heading("7.2 Readback 사용", level=2)
    add_bullet(doc, "Readback은 진단 시점에 명시적으로 수행하고 고속 animation Tick마다 호출하지 않는다.")
    add_bullet(doc, "Readback 전송 전 queue flush 및 응답 timeout이 추가되므로 제어 주기에 영향을 줄 수 있다.")

    page_break(doc)
    doc.add_heading("8. 오류 및 상태 관리", level=1)
    add_table(doc, ["조건", "대표 오류/상태", "OEM 처리"], [
        ("Bridge 미초기화", "OspBridge_Init was not called or failed", "송신 금지, 재연결 수행"),
        ("FID 미발견", "FID=0xNN not connected", "대상 remote 연결/MAC/PLCA 확인"),
        ("Multicast 미준비", "FID=0x00 multicast handle not ready", "FID별 handle 초기화 및 group 생성 확인"),
        ("FIO handle 없음", "No FID FIO OSP handle was opened", "FIO Open/Callback/RAM 상태 확인"),
        ("Queue 부족", "ADI_EAL_STATUS_NTW_BUFF_FULL", "송신 주기 완화 또는 queue/flush 정책 확인"),
        ("전송 시작 실패", "ADI_EAL_STATUS_START_FAILED", "SPI/AD3306 transport 상태 확인"),
        ("응답 timeout", "ADI_EAL_STATUS_RX_TIMEOUT", "Readback/초기화에 한해 링크 및 LED chain 확인"),
        ("미발견 FID 상태", "-99", "해당 FID가 실제 미장착이면 정상"),
        ("FIO0 skip 상태", "-77", "현재 빌드 모드와 실제 FID별 handle 상태를 함께 판정"),
    ], [2200, 3500, 3660], font_size=8.4)

    doc.add_heading("8.1 연결 로그 판정 우선순위", level=2)
    add_bullet(doc, "1순위: AD3304 FID discovery 결과 및 연결 mask")
    add_bullet(doc, "2순위: FID별 Open/CB/RAM/Reset/Bidir/Setup/State/Hdl 상태")
    add_bullet(doc, "3순위: 실제 SetColor 호출 반환값과 LastError")
    add_bullet(doc, "4순위: 필요 시 명시적 Readback(PWM/temperature)")

    doc.add_heading("9. OEM 통합 요구사항", level=1)
    add_table(doc, ["ID", "요구사항", "검증 기준"], [
        ("COM-001", "OEM SW는 FID 0x01~0x08과 nodeAddr를 분리해 관리해야 한다.", "잘못된 remote/LED 선택 없음"),
        ("COM-002", "RGB 입력 범위는 각 0~255이어야 한다.", "0/128/255 경계 PWM 일치"),
        ("COM-003", "복수 LED 동시 갱신은 batch API를 사용해야 한다.", "Tick당 flush 1회"),
        ("COM-004", "미발견 FID에는 송신하지 않아야 한다.", "not connected 오류 선검출"),
        ("COM-005", "연결 성공은 FID별 Init 상태로 판정해야 한다.", "필수 단계 status=0, Hdl=1"),
        ("COM-006", "고속 주기 제어에서 Readback을 분리해야 한다.", "animation jitter 허용치 만족"),
        ("COM-007", "오류 발생 시 LastError를 진단 로그에 보존해야 한다.", "FID/node/count 포함 로그"),
        ("COM-008", "전원 재인가 후 InitFixed 및 discovery를 다시 수행해야 한다.", "remote 재발견 및 LED 재초기화"),
    ], [1000, 5000, 3360], font_size=8.4)

    page_break(doc)
    doc.add_heading("10. OEM 검증 절차", level=1)
    add_step(doc, 1, "물리 링크 확인", "FT4222 enumeration, AD3306 SPI open, 10BASE-T1S link 및 PLCA controller 상태를 확인한다.")
    add_step(doc, 2, "Remote discovery 확인", "장착된 AD3304의 MAC 마지막 byte가 기대 FID와 일치하고 discovery mask에 반영되는지 확인한다.")
    add_step(doc, 3, "FIO 초기화 확인", "각 장착 FID의 Open/CB/RAM/Reset/Bidir/Setup/State=0, Hdl=1인지 확인한다.")
    add_step(doc, 4, "단일 LED 시험", "각 FID에서 nodeAddr 1과 마지막 node에 Red/Green/Blue/White/Off를 순차 전송한다.")
    add_step(doc, 5, "Broadcast 시험", "FID 단위 chain broadcast와 FID 0x00 전체 remote multicast의 범위를 구분하여 확인한다.")
    add_step(doc, 6, "Batch 시험", "서로 다른 RGB 배열을 4개 이상 node에 전송하고 모든 LED가 동일 Tick에 변경되는지 확인한다.")
    add_step(doc, 7, "경계값 시험", "RGB 0/1/127/128/254/255 및 nodeAddr 0/1/마지막/범위초과를 시험한다.")
    add_step(doc, 8, "성능 시험", "24/25 node animation에서 설정 주기 대비 실제 update jitter, queue full, frame loss 여부를 측정한다.")
    add_step(doc, 9, "복구 시험", "AD3304 분리/재연결, FT4222 USB 재연결, 전원 재인가 후 재초기화 동작을 확인한다.")

    doc.add_heading("10.1 권장 시험 벡터", level=2)
    add_table(doc, ["Test", "FID", "node", "RGB", "기대 결과"], [
        ("TV-01", "0x01", "1", "255,0,0", "FID1 LED1 Red full-scale"),
        ("TV-02", "0x03", "5", "255,128,0", "frame 예제와 PWM/CRC 일치"),
        ("TV-03", "0x04", "0", "0,255,0", "FID4 chain 전체 Green"),
        ("TV-04", "0x00", "0", "255,255,255", "연결된 모든 remote White"),
        ("TV-05", "미연결 FID", "1", "255,0,0", "호출 실패 + not connected"),
        ("TV-06", "연결 FID", "batch 1~N", "pattern", "모든 node 전송 후 flush 1회"),
    ], [1000, 1100, 1100, 1900, 4260], font_size=8.4)

    doc.add_heading("11. 구현 기준 파일", level=1)
    add_table(doc, ["파일", "역할"], [
        ("Forms/FormCIE.cs", "OEM/UI 입력, FID/SlaveFlag/애니메이션 데이터 생성 및 공개 C# API 호출"),
        ("FT4222/FT4222Manager.cs", "P/Invoke wrapper, 연결/오류 처리, RGB/CIE API"),
        ("NativeOsp/ex_fioOsp.c", "AD3306 network/FID discovery/FIO handle/Native public API/batch flush"),
        ("NativeOsp/utils/fioOsp.c", "FIO-OSP command framing, SET_PWM payload, CRC"),
        ("NativeOsp/utils/fioOsp.h", "PWM/Setup/Status structure 및 command ID"),
        ("NativeOsp/adi_Eth10BaseT1s_cfg.c/.h", "Controller/Remote MAC, FIO routing, E2B network configuration"),
    ], [3300, 6060], font_size=8.8)

    add_callout(doc, "문서 사용상 주의", "본 문서는 현재 프로젝트 소스 구현을 기준으로 한 OEM interface/control specification이다. AD3306/AD3304/LED IC의 전기적 timing, PHY compliance 및 공식 OSP protocol normative requirement는 해당 부품의 승인된 데이터시트와 OEM 보안 문서를 우선 적용한다.", RED)

    # Metadata
    props = doc.core_properties
    props.title = "AD3306-AD3304 LED Communication Protocol Specification"
    props.subject = "FT4222 SPI / 10BASE-T1S / Flex IO FIO-OSP OEM Interface"
    props.author = "Engineering"
    props.keywords = "AD3306, AD3304, FT4222, 10BASE-T1S, FIO, OSP, LED, OEM"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
